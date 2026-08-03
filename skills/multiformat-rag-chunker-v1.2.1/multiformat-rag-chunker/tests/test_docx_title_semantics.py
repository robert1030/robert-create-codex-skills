from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches

from helpers import FIXTURES, load_json, load_jsonl, only_source_dir, run_cli


class DocxTitleSemanticsTests(unittest.TestCase):
    def test_word_title_style_becomes_document_root(self) -> None:
        with tempfile.TemporaryDirectory() as temp_name:
            temp = Path(temp_name)
            source = temp / "title-style.docx"
            document = Document()
            document.add_paragraph("真正的文件標題", style="Title")
            document.add_heading("第一章：測試範圍", level=1)
            document.add_paragraph("測試正文。")
            document.save(source)

            code, summary = run_cli(source, temp / "output")
            self.assertEqual(code, 0, summary)
            source_dir = only_source_dir(temp / "output")
            normalized = (source_dir / "normalized-document.md").read_text(encoding="utf-8")
            ir = load_jsonl(source_dir / "document-ir.jsonl")
            chunks = load_jsonl(source_dir / "chunks.jsonl")
            report = load_json(source_dir / "processing-report.json")

            self.assertLess(normalized.index("# 真正的文件標題"), normalized.index("## 第一章：測試範圍"))
            root = next(record for record in ir if record.get("metadata", {}).get("semantic_role") == "document_title")
            chapter = next(record for record in ir if record.get("text") == "第一章：測試範圍")
            self.assertEqual(root["text"], "真正的文件標題")
            self.assertEqual(root["metadata"]["source_style_id"], "Title")
            self.assertEqual(root["metadata"]["level"], 1)
            self.assertEqual(chapter["metadata"]["level"], 2)
            self.assertEqual(chapter["heading_path"], ["真正的文件標題", "第一章：測試範圍"])
            self.assertEqual(chunks[0]["title"], "真正的文件標題")
            self.assertEqual(report["chunk_post_validation"]["document_title_mismatch_count"], 0)

    def test_inline_qr_stays_before_following_heading(self) -> None:
        with tempfile.TemporaryDirectory() as temp_name:
            temp = Path(temp_name)
            source = temp / "inline-image.docx"
            document = Document()
            document.add_paragraph("圖片順序測試", style="Title")
            document.add_paragraph("圖片前的正文。")
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(str(FIXTURES / "qr.png"), width=Inches(0.8))
            document.add_heading("下一章", level=1)
            document.add_paragraph("下一章正文。")
            document.save(source)

            code, summary = run_cli(source, temp / "output")
            self.assertEqual(code, 0, summary)
            source_dir = only_source_dir(temp / "output")
            ir = load_jsonl(source_dir / "document-ir.jsonl")
            report = load_json(source_dir / "processing-report.json")

            qr = next(record for record in ir if record.get("content_origin") == "qr_decoder")
            next_heading = next(record for record in ir if record.get("text") == "下一章")
            root = next(record for record in ir if record.get("metadata", {}).get("semantic_role") == "document_title")
            self.assertLess(qr["metadata"]["source_order"], next_heading["metadata"]["source_order"])
            self.assertEqual(qr["metadata"]["associated_heading_block_id"], root["block_id"])
            self.assertEqual(qr["heading_path"], ["圖片順序測試"])
            self.assertEqual(report["chunk_post_validation"]["reading_order_violation_count"], 0)
            self.assertEqual(report["chunk_post_validation"]["visual_heading_relation_violation_count"], 0)


if __name__ == "__main__":
    unittest.main()
