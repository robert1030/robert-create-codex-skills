from __future__ import annotations

import tempfile
import zipfile
import unittest
from pathlib import Path

from helpers import load_json, only_source_dir, run_cli


class SemanticDowngradeTests(unittest.TestCase):
    def test_multicolumn_pdf_does_not_claim_reliable_success(self) -> None:
        import fitz

        with tempfile.TemporaryDirectory() as temp_name:
            temp = Path(temp_name)
            source = temp / "two-column.pdf"
            document = fitz.open()
            page = document.new_page(width=595, height=842)
            page.insert_text((50, 50), "Two Column Test 2026-07-21", fontsize=16)
            page.insert_textbox(
                fitz.Rect(50, 100, 270, 160),
                "LEFT COLUMN ITEM ONE. This sentence must be read before left item two.",
                fontsize=11,
            )
            page.insert_textbox(
                fitz.Rect(50, 220, 270, 280),
                "LEFT COLUMN ITEM TWO. This sentence finishes the left column.",
                fontsize=11,
            )
            page.insert_textbox(
                fitz.Rect(320, 100, 545, 160),
                "RIGHT COLUMN ITEM ONE. This sentence follows the complete left column.",
                fontsize=11,
            )
            page.insert_textbox(
                fitz.Rect(320, 220, 545, 280),
                "RIGHT COLUMN ITEM TWO. This sentence finishes the right column.",
                fontsize=11,
            )
            document.save(source)
            document.close()

            code, summary = run_cli(source, temp / "output")
            self.assertEqual(code, 2, summary)
            source_dir = only_source_dir(temp / "output")
            report = load_json(source_dir / "processing-report.json")
            self.assertEqual(report["final_status"], "partial_success")
            self.assertEqual(report["source_metadata"]["layout_semantics_status"], "needs_review")
            self.assertEqual(report["source_metadata"]["multicolumn_pages"], [1])
            self.assertIn("pdf_multicolumn_layout_caution", report["warnings"])
            self.assertIn("layout_semantics_not_reliable", report["chunk_pre_validation"]["partial_reasons"])

    def test_filename_only_docx_title_is_inferred_not_reliable(self) -> None:
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as temp_name:
            temp = Path(temp_name)
            source = temp / "misleading-filename.docx"
            document = Document()
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run("Actual Human Visible Title")
            title_run.bold = True
            title_run.font.size = Pt(18)
            document.add_paragraph("First body paragraph without a Word Title or Heading style.")
            document.add_paragraph("Second body paragraph. The filename is not the document title.")
            document.save(source)

            code, summary = run_cli(source, temp / "output")
            self.assertEqual(code, 2, summary)
            source_dir = only_source_dir(temp / "output")
            report = load_json(source_dir / "processing-report.json")
            self.assertEqual(report["final_status"], "partial_success")
            self.assertEqual(report["source_metadata"]["title_source"], "filename")
            self.assertEqual(report["source_metadata"]["document_title_semantics_status"], "inferred")
            self.assertIn("document_title_inferred_from_filename", report["warnings"])
            self.assertIn(
                "document_title_semantics_not_reliable",
                report["chunk_pre_validation"]["partial_reasons"],
            )


    def test_vml_fallback_is_ignored_only_with_supported_drawingml_choice(self) -> None:
        from adapters.docx_adapter import _docx_layout_risks

        choice_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:v="urn:schemas-microsoft-com:vml"
 xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
 <w:body><w:p><mc:AlternateContent><mc:Choice Requires="a"><w:drawing><a:graphic/></w:drawing></mc:Choice>
 <mc:Fallback><w:pict><v:shape/></w:pict></mc:Fallback></mc:AlternateContent></w:p></w:body></w:document>"""
        fallback_only_xml = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
 xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
 xmlns:v="urn:schemas-microsoft-com:vml">
 <w:body><w:p><mc:AlternateContent><mc:Fallback><w:pict><v:shape/></w:pict></mc:Fallback>
 </mc:AlternateContent></w:p></w:body></w:document>"""
        with tempfile.TemporaryDirectory() as temp_name:
            temp = Path(temp_name)
            supported = temp / "supported-choice.docx"
            unsupported = temp / "fallback-only.docx"
            for path, xml in ((supported, choice_xml), (unsupported, fallback_only_xml)):
                with zipfile.ZipFile(path, "w") as archive:
                    archive.writestr("word/document.xml", xml)
            self.assertNotIn("docx_vml_layout_caution", _docx_layout_risks(supported))
            self.assertIn("docx_vml_layout_caution", _docx_layout_risks(unsupported))

if __name__ == "__main__":
    unittest.main()
