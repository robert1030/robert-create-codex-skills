#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX adapter with native OOXML, DrawingML, and rendered-PDF fallback paths."""

from __future__ import annotations

import hashlib
import posixpath
import re
import shutil
import subprocess
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from adapters.base import AdapterContext
from adapters.pdf_adapter import parse_pdf_document
from models import DocumentIR, Location
from ocr import ocr_image
from utils import normalize_nfc
from visual import classify_visual, load_image

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
WPG_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
V_NS = "urn:schemas-microsoft-com:vml"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _docx_layout_risks(path: Path) -> list[str]:
    """Detect package structures outside the supported visual-order parser path."""
    risks: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            document_xml = archive.read("word/document.xml") if "word/document.xml" in names else b""
            if document_xml:
                from lxml import etree

                root = etree.fromstring(document_xml)
                namespaces = {"w": W_NS, "v": V_NS, "mc": MC_NS}
                unsupported_vml = root.xpath(
                    "//*[self::w:pict or self::v:shape or self::v:group]"
                    "[not(ancestor::mc:Fallback)]",
                    namespaces=namespaces,
                )
                fallback_without_supported_choice = []
                for fallback in root.xpath(
                    "//mc:Fallback[.//w:pict or .//v:shape or .//v:group]",
                    namespaces=namespaces,
                ):
                    alternate_content = fallback.getparent()
                    supported_choices = alternate_content.xpath(
                        "./mc:Choice[.//w:drawing or .//a:graphic]",
                        namespaces={**namespaces, "a": A_NS},
                    ) if alternate_content is not None else []
                    if not supported_choices:
                        fallback_without_supported_choice.append(fallback)
                if unsupported_vml or fallback_without_supported_choice:
                    risks.append("docx_vml_layout_caution")
            for name in names:
                if not re.fullmatch(r"word/_rels/(?:header|footer)\d+\.xml\.rels", name):
                    continue
                if b"/image" in archive.read(name):
                    risks.append("docx_header_footer_visual_caution")
                    break
    except Exception as exc:
        risks.append(f"docx_layout_risk_scan_failed:{type(exc).__name__}")
    return risks



def _shape_text_count(path: Path) -> int:
    try:
        from lxml import etree

        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = etree.fromstring(xml)
        namespaces = {"w": W_NS}
        return len("".join(root.xpath("//w:t/text()", namespaces=namespaces)))
    except Exception:
        return 0


def _convert_to_pdf(path: Path, work_dir: Path) -> Path:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise RuntimeError("libreoffice_not_available")
    output_dir = work_dir / "docx-rendered-pdf"
    output_dir.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run(
        [executable, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(path)],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        check=False,
        timeout=120,
    )
    candidate = output_dir / f"{path.stem}.pdf"
    if completed.returncode != 0 or not candidate.is_file() or candidate.stat().st_size == 0:
        raise RuntimeError(f"docx_to_pdf_failed:{completed.stdout}:{completed.stderr}")
    return candidate


def _iter_block_items(document):
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn

    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _relationship_targets(path: Path) -> dict[str, str]:
    """Return main-document image relationship targets as package paths."""
    from lxml import etree

    rels_path = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(path) as archive:
        if rels_path not in archive.namelist():
            return {}
        root = etree.fromstring(archive.read(rels_path))
    targets: dict[str, str] = {}
    for relationship in root.findall(f"{{{REL_NS}}}Relationship"):
        if relationship.get("TargetMode") == "External":
            continue
        rel_id = relationship.get("Id")
        target = relationship.get("Target")
        rel_type = relationship.get("Type", "")
        if not rel_id or not target or not rel_type.endswith("/image"):
            continue
        targets[rel_id] = posixpath.normpath(posixpath.join("word", target))
    return targets


def _element_image_occurrences(
    element,
    relationship_targets: dict[str, str],
    *,
    element_index: int,
    page: int | None = None,
) -> list[dict[str, Any]]:
    """Extract image occurrences from a paragraph or table in XML order."""
    from lxml import etree

    namespaces = {"wp": WP_NS, "a": A_NS, "r": R_NS}
    root = etree.fromstring(etree.tostring(element))
    occurrences: list[dict[str, Any]] = []
    for occurrence_index, blip in enumerate(root.xpath(".//a:blip[@r:embed]", namespaces=namespaces), start=1):
        rel_id = blip.get(f"{{{R_NS}}}embed")
        if not rel_id:
            continue
        drawing = blip
        while drawing is not None and drawing.tag not in {f"{{{WP_NS}}}inline", f"{{{WP_NS}}}anchor"}:
            drawing = drawing.getparent()
        anchor_kind = drawing.tag.rsplit("}", 1)[-1] if drawing is not None else "unknown"
        bbox: list[float] | None = None
        if drawing is not None and anchor_kind == "anchor":
            x = _anchor_value(drawing, "./wp:positionH/wp:posOffset/text()", namespaces)
            y = _anchor_value(drawing, "./wp:positionV/wp:posOffset/text()", namespaces)
            extent = drawing.xpath("./wp:extent", namespaces=namespaces)
            width = int(extent[0].get("cx", "0")) if extent else 0
            height = int(extent[0].get("cy", "0")) if extent else 0
            bbox = [float(x), float(y), float(x + width), float(y + height)]
        occurrences.append({
            "relationship_id": rel_id,
            "package_path": relationship_targets.get(rel_id),
            "element_index": element_index,
            "occurrence_index": occurrence_index,
            "page": page,
            "bbox": bbox,
            "anchor_kind": anchor_kind,
        })
    return occurrences


def _media_block_from_occurrence(
    archive: zipfile.ZipFile,
    context: AdapterContext,
    occurrence: dict[str, Any],
    *,
    asset_id: str,
    heading_path: list[str],
    seen_payloads: set[str],
):
    package_path = occurrence.get("package_path")
    location = Location(
        page=occurrence.get("page"),
        bbox=occurrence.get("bbox"),
        element_index=occurrence.get("element_index"),
        asset_id=asset_id,
    )
    provenance_metadata = {
        "asset_id": asset_id,
        "package_path": package_path,
        "relationship_id": occurrence.get("relationship_id"),
        "occurrence_index": occurrence.get("occurrence_index"),
        "anchor_kind": occurrence.get("anchor_kind"),
        "association_method": occurrence.get("association_method", "ooxml_container_order"),
    }
    if occurrence.get("original_bbox") is not None:
        provenance_metadata["original_bbox"] = occurrence["original_bbox"]
    try:
        if not package_path or package_path not in archive.namelist():
            raise FileNotFoundError(f"missing_related_media:{package_path}")
        image_bytes = archive.read(package_path)
        asset_sha256 = hashlib.sha256(image_bytes).hexdigest()
        image = load_image(image_bytes)
        inspection = classify_visual(image, name_hint=Path(package_path).name)
        visual_review = (
            context.visual_semantics.lookup(str(package_path), asset_sha256)
            if context.visual_semantics is not None
            else None
        )
        metadata = {
            **provenance_metadata,
            "visual_class": inspection.visual_class,
            "width": inspection.width,
            "height": inspection.height,
            "asset_sha256": asset_sha256,
            "machine_payloads": [],
        }
        unique_qr = [value for value in inspection.qr_payloads if value not in seen_payloads]
        if unique_qr:
            seen_payloads.update(unique_qr)
            return context.block(
                "image",
                "\n".join(f"QR Code payload：{value}" for value in unique_qr),
                location=location,
                heading_path=list(heading_path),
                content_origin="qr_decoder",
                required=True,
                critical=True,
                metadata={**metadata, "qr_payloads": unique_qr},
            )
        if inspection.qr_payloads:
            return context.block(
                "image", "", location=location, heading_path=list(heading_path),
                content_origin="derived_normalization", required=False, status="skipped", verbatim=False,
                metadata={**metadata, "skip_reason": "duplicate_qr_payload", "qr_payloads": inspection.qr_payloads},
            )
        if visual_review is not None and visual_review.review_mode == "semantic_summary":
            route, _admission = context.visual_route(
                str(package_path), asset_sha256, review_present=True,
                visual_class=inspection.visual_class, native_structured_parser_status="insufficient",
            )
            return context.block(
                "image", visual_review.summary, location=location, heading_path=list(heading_path),
                content_origin="llm_visual_summary", required=True, verbatim=False,
                metadata={
                    **metadata,
                    "visual_summary_evidence": visual_review.evidence(),
                    "capability_route": route,
                },
            )
        if inspection.visual_class in {"logo", "icon", "photo", "decorative", "screen_capture"}:
            route = None
            if inspection.visual_class == "screen_capture":
                route, _admission = context.visual_route(
                    str(package_path), asset_sha256, review_present=False,
                    visual_class=inspection.visual_class, native_structured_parser_status="insufficient",
                )
            return context.block(
                "image", "", location=location, heading_path=list(heading_path),
                content_origin="derived_normalization", required=False, status="skipped", verbatim=False,
                metadata={
                    **metadata,
                    "skip_reason": "non_required_visual",
                    **({"capability_route": route} if route is not None else {}),
                },
            )
        route, admission = context.visual_route(
            str(package_path), asset_sha256, review_present=False,
            visual_class=inspection.visual_class, native_structured_parser_status="insufficient",
        )
        result = ocr_image(image_bytes, admission=admission, languages=context.ocr_languages)
        if result.text and inspection.width <= 400 and inspection.height <= 400 and len(result.text.strip()) < 30:
            return context.block(
                "image", "", location=location, heading_path=list(heading_path),
                content_origin="derived_normalization", required=False, status="skipped", verbatim=False,
                metadata={
                    **metadata,
                    "skip_reason": "small_logo_or_icon_candidate",
                    "ocr_candidate_rejected": result.text,
                    "capability_route": route,
                },
            )
        status = "success" if result.status == "success" else ("low_quality" if result.text else "failed")
        block = context.block(
            "image",
            result.text if result.status == "success" else "",
            location=location,
            heading_path=list(heading_path),
            content_origin="ocr" if result.text else "placeholder",
            required=True,
            status=status,
            metadata={
                **metadata,
                "ocr_confidence": result.confidence,
                "ocr_quality": result.quality,
                "ocr_semantic_status": "accepted" if result.status == "success" else "not_run" if result.status == "blocked" else "rejected",
                "reason": result.quality.get("reasons", [None])[0] if result.status != "success" else None,
                "capability_route": route,
            },
        )
        block.attempts = result.attempts
        return block
    except Exception as exc:
        return context.block(
            "image", "", location=location, heading_path=list(heading_path),
            content_origin="placeholder", required=True, critical=False, status="failed", verbatim=False,
            metadata={**provenance_metadata, "error": str(exc)},
        )


def _append_media_occurrences(
    path: Path,
    context: AdapterContext,
    blocks: list,
    occurrences: list[dict[str, Any]],
    *,
    heading_path: list[str],
    seen_payloads: set[str],
    starting_index: int,
) -> int:
    image_index = starting_index
    if not occurrences:
        return image_index
    with zipfile.ZipFile(path) as archive:
        for occurrence in occurrences:
            image_index += 1
            blocks.append(_media_block_from_occurrence(
                archive,
                context,
                occurrence,
                asset_id=f"docx-image-{image_index:03d}",
                heading_path=heading_path,
                seen_payloads=seen_payloads,
            ))
    return image_index


def _paragraph_style(paragraph) -> tuple[str, str]:
    try:
        style = paragraph.style
        return str(style.name or ""), str(style.style_id or "")
    except Exception:
        return "", ""


def _is_title_style(style_name: str, style_id: str) -> bool:
    normalized = {style_name.strip().lower(), style_id.strip().lower()}
    return bool(normalized & {"title", "document title"})


def _heading_level(style_name: str, style_id: str) -> int | None:
    for value in (style_id, style_name):
        match = re.search(r"heading\s*([1-9])", value, flags=re.IGNORECASE)
        if match:
            return max(1, min(6, int(match.group(1))))
    return None


def _annotate_reading_order(blocks: list, document_title: str) -> None:
    """Assign deterministic source order and nearest-heading relationships."""
    heading_text: list[str] = []
    heading_ids: list[str | None] = []
    for source_order, block in enumerate(blocks, start=1):
        block.metadata["source_order"] = source_order
        if block.location.element_index is None:
            block.location.element_index = source_order
        if block.type == "heading" and block.status == "success":
            level = max(1, min(6, int(block.metadata.get("level", 1))))
            if level == 1:
                heading_text = [block.text]
                heading_ids = [block.block_id]
            else:
                if not heading_text and document_title:
                    heading_text = [document_title]
                    heading_ids = [None]
                heading_text = heading_text[: level - 1]
                heading_ids = heading_ids[: level - 1]
                while len(heading_text) < level - 1:
                    heading_text.append(document_title)
                    heading_ids.append(None)
                heading_text.append(block.text)
                heading_ids.append(block.block_id)
            block.heading_path = list(heading_text)
            parent_id = next((value for value in reversed(heading_ids[:-1]) if value), None)
            if parent_id:
                block.metadata["parent_heading_block_id"] = parent_id
            continue
        block.heading_path = list(heading_text or ([document_title] if document_title else []))
        nearest_heading_id = next((value for value in reversed(heading_ids) if value), None)
        if nearest_heading_id:
            block.metadata["associated_heading_block_id"] = nearest_heading_id
            block.metadata["associated_heading_path"] = list(block.heading_path)
        if block.type == "image":
            block.metadata.setdefault("association_method", "nearest_preceding_heading")


def _native_text_length(path: Path) -> int:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(path)
    total = 0
    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            total += len(normalize_nfc(item.text))
        elif isinstance(item, Table):
            total += sum(len(normalize_nfc(cell.text)) for row in item.rows for cell in row.cells)
    return total


def _native_parse(path: Path, context: AdapterContext) -> DocumentIR:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(path)
    relationship_targets = _relationship_targets(path)
    explicit_title = next(
        (
            normalize_nfc(paragraph.text)
            for paragraph in document.paragraphs
            if normalize_nfc(paragraph.text)
            and _is_title_style(*_paragraph_style(paragraph))
        ),
        "",
    )
    core_title = normalize_nfc(str(document.core_properties.title or ""))
    title = explicit_title or core_title or path.stem
    title_source = "title_style" if explicit_title else ("core_properties" if core_title else "pending_heading_or_filename")
    root_offset = 1 if explicit_title or core_title else 0
    title_role_assigned = False
    blocks: list = []
    element_index = 0
    image_index = 0
    seen_payloads: set[str] = set()

    if core_title and not explicit_title:
        blocks.append(context.block(
            "heading", core_title,
            location=Location(element_index=0),
            heading_path=[core_title],
            content_origin="derived_normalization",
            required=True,
            critical=True,
            verbatim=False,
            metadata={
                "level": 1,
                "semantic_role": "document_title",
                "title_source": "core_properties",
            },
        ))
        title_role_assigned = True

    for item in _iter_block_items(document):
        element_index += 1
        if isinstance(item, Paragraph):
            text = normalize_nfc(item.text)
            style_name, style_id = _paragraph_style(item)
            occurrences = _element_image_occurrences(
                item._p,
                relationship_targets,
                element_index=element_index,
            )
            if text:
                if _is_title_style(style_name, style_id):
                    title = text
                    title_source = "title_style"
                    title_role_assigned = True
                    blocks.append(context.block(
                        "heading", text,
                        location=Location(element_index=element_index),
                        heading_path=[text],
                        content_origin="native_text",
                        required=True,
                        critical=True,
                        metadata={
                            "level": 1,
                            "style": style_name,
                            "source_style_name": style_name,
                            "source_style_id": style_id,
                            "semantic_role": "document_title",
                            "title_source": "title_style",
                        },
                    ))
                else:
                    raw_heading_level = _heading_level(style_name, style_id)
                    if raw_heading_level is not None:
                        level = min(6, raw_heading_level + root_offset)
                        semantic_role = "section_heading"
                        if raw_heading_level == 1 and not title_role_assigned and root_offset == 0:
                            title = text
                            title_source = "heading_1"
                            title_role_assigned = True
                            semantic_role = "document_title"
                        blocks.append(context.block(
                            "heading", text,
                            location=Location(element_index=element_index),
                            heading_path=[],
                            content_origin="native_text",
                            required=True,
                            critical=True,
                            metadata={
                                "level": level,
                                "source_heading_level": raw_heading_level,
                                "style": style_name,
                                "source_style_name": style_name,
                                "source_style_id": style_id,
                                "semantic_role": semantic_role,
                                **({"title_source": "heading_1"} if semantic_role == "document_title" else {}),
                            },
                        ))
                    elif style_name.lower().startswith("list"):
                        blocks.append(context.block(
                            "list", f"- {text}",
                            location=Location(element_index=element_index),
                            heading_path=[],
                            content_origin="native_text",
                            required=True,
                            metadata={"ordered": False, "items": [text], "style": style_name, "source_style_id": style_id},
                        ))
                    else:
                        critical = bool(re.search(r"\b20\d{2}[-/.]\d{1,2}[-/.]\d{1,2}\b|\bEP\.?\s*\d+\b", text, flags=re.IGNORECASE))
                        blocks.append(context.block(
                            "paragraph", text,
                            location=Location(element_index=element_index),
                            heading_path=[],
                            content_origin="native_text",
                            required=True,
                            critical=critical,
                            metadata={"style": style_name, "source_style_id": style_id},
                        ))
            image_index = _append_media_occurrences(
                path,
                context,
                blocks,
                occurrences,
                heading_path=[],
                seen_payloads=seen_payloads,
                starting_index=image_index,
            )
        elif isinstance(item, Table):
            rows = [[normalize_nfc(cell.text) for cell in row.cells] for row in item.rows]
            if rows:
                header = rows[0]
                data = rows[1:]
                blocks.append(context.block(
                    "table", "",
                    location=Location(element_index=element_index),
                    heading_path=[],
                    content_origin="native_table",
                    required=True,
                    critical=True,
                    metadata={
                        "header": header,
                        "rows": data,
                        "logical_column_count": len(header),
                        "data_row_count": len(data),
                    },
                ))
            occurrences = _element_image_occurrences(
                item._tbl,
                relationship_targets,
                element_index=element_index,
            )
            for occurrence in occurrences:
                occurrence["association_method"] = "table_container_order"
            image_index = _append_media_occurrences(
                path,
                context,
                blocks,
                occurrences,
                heading_path=[],
                seen_payloads=seen_payloads,
                starting_index=image_index,
            )

    if not title_role_assigned:
        blocks.insert(0, context.block(
            "heading", title,
            location=Location(element_index=0),
            heading_path=[title],
            content_origin="derived_normalization",
            required=True,
            critical=True,
            verbatim=False,
            metadata={
                "level": 1,
                "semantic_role": "document_title",
                "title_source": "filename",
                "title_from_filename": True,
            },
        ))
        title_source = "filename"
    _annotate_reading_order(blocks, title)
    layout_risks = _docx_layout_risks(path)
    title_semantics_status = "reliable" if title_source != "filename" else "inferred"
    warnings = list(layout_risks)
    if title_semantics_status != "reliable":
        warnings.append("document_title_inferred_from_filename")
    return DocumentIR(
        source_id=context.provenance.source_id,
        title=title,
        provenance=context.provenance,
        blocks=blocks,
        metadata={
            "adapter": "docx_adapter",
            "native_element_count": element_index,
            "visual_occurrence_count": image_index,
            "document_title": title,
            "title_source": title_source,
            "layout_semantics_status": "needs_review" if layout_risks else "reliable",
            "document_title_semantics_status": title_semantics_status,
        },
        warnings=warnings,
    )


def _anchor_value(anchor, xpath: str, namespaces: dict[str, str], default: int = 0) -> int:
    values = anchor.xpath(xpath, namespaces=namespaces)
    try:
        return int(values[0]) if values else default
    except (TypeError, ValueError):
        return default


def _cluster_positions(values: list[int], tolerance: int = 20000) -> list[tuple[int, int]]:
    if not values:
        return []
    clusters: list[list[int]] = []
    for value in sorted(values):
        if not clusters or abs(value - int(sum(clusters[-1]) / len(clusters[-1]))) > tolerance:
            clusters.append([value])
        else:
            clusters[-1].append(value)
    return [(int(round(sum(cluster) / len(cluster))), len(cluster)) for cluster in clusters]


def _clean_shape_text(text: str) -> str:
    text = normalize_nfc(text)
    text = re.sub(r"[ \t]+", " ", text).strip()
    text = re.sub(r"\s+([，。；：！？、,.!?;:])", r"\1", text)
    return text


def _xfrm_geometry(xfrm) -> tuple[float, float, float, float, float, float, float, float]:
    namespaces = {"a": A_NS}
    off = xfrm.xpath("./a:off", namespaces=namespaces)
    ext = xfrm.xpath("./a:ext", namespaces=namespaces)
    child_off = xfrm.xpath("./a:chOff", namespaces=namespaces)
    child_ext = xfrm.xpath("./a:chExt", namespaces=namespaces)
    x = float(off[0].get("x", "0")) if off else 0.0
    y = float(off[0].get("y", "0")) if off else 0.0
    width = float(ext[0].get("cx", "0")) if ext else 0.0
    height = float(ext[0].get("cy", "0")) if ext else 0.0
    child_x = float(child_off[0].get("x", "0")) if child_off else 0.0
    child_y = float(child_off[0].get("y", "0")) if child_off else 0.0
    child_width = float(child_ext[0].get("cx", str(width or 1.0))) if child_ext else (width or 1.0)
    child_height = float(child_ext[0].get("cy", str(height or 1.0))) if child_ext else (height or 1.0)
    return x, y, width, height, child_x, child_y, child_width, child_height


def _node_xfrm(node):
    namespaces = {"a": A_NS, "pic": PIC_NS, "wpg": WPG_NS, "wps": WPS_NS}
    local = node.tag.rsplit("}", 1)[-1]
    if local == "pic":
        values = node.xpath("./pic:spPr/a:xfrm", namespaces=namespaces)
    elif local in {"wgp", "grpSp"}:
        values = node.xpath("./wpg:grpSpPr/a:xfrm|./a:grpSpPr/a:xfrm", namespaces=namespaces)
    else:
        values = node.xpath("./wps:spPr/a:xfrm|./a:xfrm", namespaces=namespaces)
    return values[0] if values else None


def _picture_bbox_in_anchor(anchor, blip, *, anchor_x: int, anchor_y: int, anchor_width: int, anchor_height: int) -> list[float]:
    picture = blip
    while picture is not None and picture is not anchor and picture.tag.rsplit("}", 1)[-1] != "pic":
        picture = picture.getparent()
    if picture is None or picture is anchor:
        return [float(anchor_x), float(anchor_y), float(anchor_x + anchor_width), float(anchor_y + anchor_height)]
    picture_xfrm = _node_xfrm(picture)
    if picture_xfrm is None:
        return [float(anchor_x), float(anchor_y), float(anchor_x + anchor_width), float(anchor_y + anchor_height)]
    x, y, width, height, _child_x, _child_y, _child_width, _child_height = _xfrm_geometry(picture_xfrm)
    parent = picture.getparent()
    while parent is not None and parent is not anchor:
        parent_xfrm = _node_xfrm(parent)
        if parent_xfrm is not None:
            off_x, off_y, ext_w, ext_h, child_x, child_y, child_w, child_h = _xfrm_geometry(parent_xfrm)
            scale_x = ext_w / child_w if child_w else 1.0
            scale_y = ext_h / child_h if child_h else 1.0
            x = off_x + (x - child_x) * scale_x
            y = off_y + (y - child_y) * scale_y
            width *= scale_x
            height *= scale_y
        parent = parent.getparent()
    if width <= 0 or height <= 0:
        width = float(anchor_width)
        height = float(anchor_height)
    return [
        float(anchor_x) + x,
        float(anchor_y) + y,
        float(anchor_x) + x + width,
        float(anchor_y) + y + height,
    ]


def _extract_shape_pages(path: Path) -> tuple[list[list[dict[str, Any]]], int]:
    from lxml import etree

    relationship_targets = _relationship_targets(path)
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = etree.fromstring(xml)
    namespaces = {
        "w": W_NS,
        "wp": WP_NS,
        "a": A_NS,
        "r": R_NS,
        "pic": PIC_NS,
        "wpg": WPG_NS,
        "wps": WPS_NS,
    }
    pages: list[list[dict[str, Any]]] = []
    for paragraph in root.xpath("//w:body/w:p", namespaces=namespaces):
        anchors = paragraph.xpath(".//wp:anchor", namespaces=namespaces)
        if not anchors:
            continue
        page_items: list[dict[str, Any]] = []
        for anchor_index, anchor in enumerate(anchors, start=1):
            text = _clean_shape_text("".join(anchor.xpath(".//w:t/text()", namespaces=namespaces)))
            x = _anchor_value(anchor, "./wp:positionH/wp:posOffset/text()", namespaces)
            y = _anchor_value(anchor, "./wp:positionV/wp:posOffset/text()", namespaces)
            extent = anchor.xpath("./wp:extent", namespaces=namespaces)
            width = int(extent[0].get("cx", "0")) if extent else 0
            height = int(extent[0].get("cy", "0")) if extent else 0
            image_occurrences: list[dict[str, Any]] = []
            for occurrence_index, blip in enumerate(anchor.xpath(".//a:blip[@r:embed]", namespaces=namespaces), start=1):
                rel_id = blip.get(f"{{{R_NS}}}embed")
                if not rel_id:
                    continue
                bbox = _picture_bbox_in_anchor(
                    anchor,
                    blip,
                    anchor_x=x,
                    anchor_y=y,
                    anchor_width=width,
                    anchor_height=height,
                )
                image_occurrences.append({
                    "relationship_id": rel_id,
                    "package_path": relationship_targets.get(rel_id),
                    "occurrence_index": occurrence_index,
                    "anchor_index": anchor_index,
                    "anchor_kind": "drawingml_group_or_anchor",
                    "bbox": bbox,
                    "original_bbox": list(bbox),
                    "association_method": "drawingml_geometry",
                })
            page_items.append({
                "anchor_index": anchor_index,
                "text": text,
                "x": x,
                "y": y,
                "width": width,
                "height": height,
                "embeds": [value["relationship_id"] for value in image_occurrences],
                "image_occurrences": image_occurrences,
            })
        pages.append(page_items)
    return pages, len(pages)


def _longest_numeric_sequence(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    numeric = [item for item in items if re.fullmatch(r"\d+", item["text"])]
    numeric.sort(key=lambda item: (item["y"], item["x"]))
    best: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []
    previous: int | None = None
    for item in numeric:
        value = int(item["text"])
        if previous is None or value == previous + 1:
            current.append(item)
        else:
            if len(current) > len(best):
                best = current
            current = [item]
        previous = value
    if len(current) > len(best):
        best = current
    return best


def _shape_table(items: list[dict[str, Any]]) -> tuple[list[list[str]], set[int]]:
    sequence = _longest_numeric_sequence(items)
    if len(sequence) < 3:
        return [], set()
    row_cells: list[tuple[dict[str, Any], list[dict[str, Any]]]] = []
    all_cell_x: list[int] = []
    for number_item in sequence:
        tolerance = max(2500, int(number_item["height"] * 0.08))
        same_row = [
            item for item in items
            if item["text"] and item["anchor_index"] != number_item["anchor_index"]
            and abs(item["y"] - number_item["y"]) <= tolerance
            and item["x"] > number_item["x"]
        ]
        same_row.sort(key=lambda item: item["x"])
        row_cells.append((number_item, same_row))
        all_cell_x.extend(item["x"] for item in same_row)
    clusters = _cluster_positions(all_cell_x)
    frequent = [center for center, count in clusters if count >= max(2, int(len(sequence) * 0.4))]
    if len(frequent) < 2:
        frequent = [center for center, _count in sorted(clusters, key=lambda pair: pair[1], reverse=True)[:3]]
    frequent = sorted(frequent)[:3]
    if len(frequent) != 3:
        return [], set()
    rows: list[list[str]] = []
    used: set[int] = set()
    for number_item, cells in row_cells:
        row = [number_item["text"], "", "", ""]
        used.add(number_item["anchor_index"])
        for cell in cells:
            target = min(range(3), key=lambda index: abs(cell["x"] - frequent[index])) + 1
            if row[target]:
                row[target] = f"{row[target]} {cell['text']}".strip()
            else:
                row[target] = cell["text"]
            used.add(cell["anchor_index"])
        row[3] = re.sub(
            r"\s*/\s*(?=(?:adj|adv|n|v|prep|conj|pron)\.)",
            " / ",
            row[3],
            flags=re.IGNORECASE,
        )
        rows.append(row)
    expected = list(range(int(rows[0][0]), int(rows[0][0]) + len(rows)))
    if [int(row[0]) for row in rows] != expected:
        return [], set()
    return rows, used


def _shape_numbered_items(items: list[dict[str, Any]], excluded: set[int]) -> tuple[list[str], set[int]]:
    candidates = [
        item for item in sorted(items, key=lambda value: (value["y"], value["x"]))
        if item["text"] and item["anchor_index"] not in excluded
        and not _is_shape_page_label(item["text"])
    ]
    result: list[str] = []
    used: set[int] = set()
    current_index: int | None = None
    current_parts: list[str] = []
    for item in candidates:
        match = re.match(r"^(\d+)[.)]\s*(.*)$", item["text"])
        if match:
            if current_index is not None:
                result.append(" ".join(current_parts).strip())
            current_index = int(match.group(1))
            current_parts = [match.group(2).strip()]
            used.add(item["anchor_index"])
        elif current_index is not None:
            current_parts.append(item["text"])
            used.add(item["anchor_index"])
    if current_index is not None:
        result.append(" ".join(current_parts).strip())
    if len(result) < 2:
        return [], set()
    return result, used


def _items_bbox(items: list[dict[str, Any]]) -> list[float] | None:
    if not items:
        return None
    return [
        float(min(item["x"] for item in items)),
        float(min(item["y"] for item in items)),
        float(max(item["x"] + item["width"] for item in items)),
        float(max(item["y"] + item["height"] for item in items)),
    ]


def _shape_text_key(text: str) -> str:
    return re.sub(r"\s+", "", text).upper()


def _is_shape_page_label(text: str) -> bool:
    return bool(re.fullmatch(r"(?:PAGE|PAG\s*E?)?\s*\d+", text.strip(), re.IGNORECASE))


def _shape_article_title_candidate(
    items: list[dict[str, Any]],
    title_item: dict[str, Any],
    excluded: set[int],
    structured_indices: set[int],
) -> dict[str, Any] | None:
    page_width = max((int(item["x"]) + int(item["width"]) for item in items), default=1)
    title_bottom = int(title_item["y"]) + int(title_item["height"])
    structured_top = min(
        (int(item["y"]) for item in items if item["anchor_index"] in structured_indices),
        default=10**18,
    )
    candidates = [
        item for item in items
        if item["text"]
        and item["anchor_index"] not in excluded
        and title_bottom <= int(item["y"]) < structured_top
        and 8 <= len(item["text"].strip()) <= 220
        and int(item["width"]) >= page_width * 0.3
        and not _is_shape_page_label(item["text"])
    ]
    return min(candidates, key=lambda item: (item["y"], item["x"]), default=None)


def _nearest_preceding_shape_label(
    items: list[dict[str, Any]],
    structured_indices: set[int],
    excluded: set[int],
) -> dict[str, Any] | None:
    structured_top = min(
        (int(item["y"]) for item in items if item["anchor_index"] in structured_indices),
        default=None,
    )
    if structured_top is None:
        return None
    candidates = [
        item for item in items
        if item["text"]
        and item["anchor_index"] not in structured_indices
        and item["anchor_index"] not in excluded
        and int(item["y"]) + int(item["height"]) <= structured_top
        and len(item["text"].strip()) <= 120
        and not _is_shape_page_label(item["text"])
    ]
    return max(candidates, key=lambda item: (item["y"] + item["height"], item["x"]), default=None)


def _numbered_section_title(items: list[str]) -> str:
    joined = " ".join(items)
    source_characters = len(re.findall(r"\S", joined))
    latin_characters = len(re.findall(r"[A-Za-z]", joined))
    return "英文段落" if source_characters and latin_characters / source_characters >= 0.6 else "編號段落"


def _drawingml_sort_blocks(blocks: list) -> list:
    body_start_by_page: dict[int, float] = {}
    header_bottom_by_page: dict[int, float] = {}
    for block in blocks:
        page = block.location.page
        bbox = block.location.bbox
        if page is None or not bbox:
            continue
        semantic_role = block.metadata.get("semantic_role")
        is_body_start = block.type in {"table", "list"} or semantic_role == "section_heading"
        if is_body_start:
            body_start_by_page[page] = min(body_start_by_page.get(page, float("inf")), float(bbox[1]))
    for block in blocks:
        page = block.location.page
        bbox = block.location.bbox
        if page is None or not bbox:
            continue
        if block.type != "heading" or block.metadata.get("semantic_role") not in {"document_title", "article_title"}:
            continue
        body_start = body_start_by_page.get(page, float("inf"))
        if float(bbox[1]) < body_start:
            header_bottom_by_page[page] = max(header_bottom_by_page.get(page, 0.0), float(bbox[3]))

    creation_order = {block.block_id: index for index, block in enumerate(blocks)}

    def sort_key(block) -> tuple[float, float, float, int, int]:
        page = float(block.location.page or 0)
        bbox = block.location.bbox
        y = float(bbox[1]) if bbox else float(block.location.element_index or 10**12)
        x = float(bbox[0]) if bbox else 0.0
        if block.type == "image" and bbox and block.location.page is not None:
            body_start = body_start_by_page.get(block.location.page)
            header_bottom = header_bottom_by_page.get(block.location.page)
            if body_start is not None and header_bottom is not None and y < body_start:
                block.metadata["reading_order_adjustment"] = "header_visual_after_heading_cluster"
                block.metadata["association_method"] = "drawingml_header_cluster"
                block.metadata["original_reading_y"] = y
                y = header_bottom + 0.5 + float(block.metadata.get("occurrence_index", 0)) / 1000.0
        priority = 0 if block.type == "heading" else (1 if block.type == "image" else 2)
        return page, y, x, priority, creation_order[block.block_id]

    return sorted(blocks, key=sort_key)


def _drawingml_parse(path: Path, context: AdapterContext) -> DocumentIR:
    pages, page_count = _extract_shape_pages(path)
    if not pages:
        raise RuntimeError("no_drawingml_pages")
    context.provenance.source_dimensions["page_count"] = page_count
    blocks: list = []
    title = path.stem
    title_source = "filename"
    used_by_page: dict[int, set[int]] = {index: set() for index in range(1, page_count + 1)}
    shaped_tables = {
        page_number: _shape_table(items)
        for page_number, items in enumerate(pages, start=1)
    }
    text_pages: dict[str, set[int]] = {}
    for page_number, items in enumerate(pages, start=1):
        for item in items:
            key = _shape_text_key(item["text"])
            if key:
                text_pages.setdefault(key, set()).add(page_number)
    repeated_header_keys = {key for key, page_numbers in text_pages.items() if len(page_numbers) >= 2}

    title_found = False
    title_position: tuple[int, dict[str, Any]] | None = None
    for page_number, items in enumerate(pages, start=1):
        for item in sorted(items, key=lambda value: (value["y"], value["x"])):
            text = item["text"]
            if text and "EP." in text and re.search(r"20\d{2}-\d{2}-\d{2}", text):
                title = text
                title_source = "drawingml_episode_heading"
                blocks.append(context.block(
                    "heading", text,
                    location=Location(page=page_number, bbox=[item["x"], item["y"], item["x"] + item["width"], item["y"] + item["height"]]),
                    heading_path=[],
                    content_origin="native_text",
                    required=True,
                    critical=True,
                    metadata={
                        "level": 1,
                        "semantic_role": "document_title",
                        "title_source": title_source,
                        "drawingml_anchor_index": item["anchor_index"],
                        "contains_date": True,
                        "contains_identifier": True,
                    },
                ))
                used_by_page[page_number].add(item["anchor_index"])
                title_position = (page_number, item)
                title_found = True
                break
        if title_found:
            break

    if title_position is not None:
        page_number, title_item = title_position
        page_items = pages[page_number - 1]
        _rows, structured_indices = shaped_tables[page_number]
        article_item = _shape_article_title_candidate(
            page_items,
            title_item,
            used_by_page[page_number],
            structured_indices,
        )
        if article_item is not None:
            blocks.append(context.block(
                "heading", article_item["text"],
                location=Location(
                    page=page_number,
                    bbox=[
                        article_item["x"], article_item["y"],
                        article_item["x"] + article_item["width"],
                        article_item["y"] + article_item["height"],
                    ],
                ),
                heading_path=[],
                content_origin="native_text",
                required=True,
                critical=True,
                metadata={
                    "level": 2,
                    "semantic_role": "article_title",
                    "drawingml_anchor_index": article_item["anchor_index"],
                    "inference": "source_order_between_document_title_and_structured_section",
                },
            ))
            used_by_page[page_number].add(article_item["anchor_index"])

    for page_number, items in enumerate(pages, start=1):
        rows, used_table = shaped_tables[page_number]
        if rows:
            section_title = _nearest_preceding_shape_label(items, used_table, used_by_page[page_number])
            section_text = section_title["text"] if section_title else f"表格 {page_number}"
            blocks.append(context.block(
                "heading", section_text,
                location=Location(
                    page=page_number,
                    bbox=[section_title["x"], section_title["y"], section_title["x"] + section_title["width"], section_title["y"] + section_title["height"]] if section_title else None,
                ),
                heading_path=[],
                content_origin="native_text",
                required=True,
                critical=True,
                metadata={
                    "level": 2,
                    "semantic_role": "section_heading",
                    "drawingml_anchor_index": section_title["anchor_index"] if section_title else None,
                },
            ))
            if section_title:
                used_by_page[page_number].add(section_title["anchor_index"])
            used_by_page[page_number].update(used_table)
            header = [f"欄位 {index}" for index in range(1, len(rows[0]) + 1)]
            used_items = [item for item in items if item["anchor_index"] in used_table]
            blocks.append(context.block(
                "table", "",
                location=Location(page=page_number, bbox=_items_bbox(used_items)),
                heading_path=[],
                content_origin="native_table",
                required=True,
                critical=True,
                metadata={
                    "header": header,
                    "rows": rows,
                    "logical_column_count": len(header),
                    "data_row_count": len(rows),
                    "source_unit_count": len(used_table),
                    "drawingml_anchor_indices": sorted(used_table),
                },
            ))

        numbered, used_numbered = _shape_numbered_items(items, used_by_page[page_number])
        if numbered:
            numbered_items = [item for item in items if item["anchor_index"] in used_numbered]
            list_bbox = _items_bbox(numbered_items)
            heading_bbox = None
            if list_bbox:
                heading_bbox = [list_bbox[0], max(0.0, list_bbox[1] - 1.0), list_bbox[2], list_bbox[1]]
            blocks.append(context.block(
                "heading", _numbered_section_title(numbered),
                location=Location(page=page_number, bbox=heading_bbox),
                heading_path=[],
                content_origin="derived_normalization",
                required=True,
                critical=False,
                verbatim=False,
                metadata={
                    "level": 2,
                    "semantic_role": "section_heading",
                    "derived": True,
                    "derivation_type": "synthetic_section_heading",
                    "derived_from_numbered_drawingml_text": True,
                },
            ))
            blocks.append(context.block(
                "list", "\n".join(f"{index}. {value}" for index, value in enumerate(numbered, start=1)),
                location=Location(page=page_number, bbox=list_bbox),
                heading_path=[],
                content_origin="native_text",
                required=True,
                critical=True,
                metadata={
                    "ordered": True,
                    "items": numbered,
                    "item_count": len(numbered),
                    "drawingml_anchor_indices": sorted(used_numbered),
                },
            ))
            used_by_page[page_number].update(used_numbered)

        for item in items:
            if not item["text"] or item["anchor_index"] in used_by_page[page_number]:
                continue
            key = _shape_text_key(item["text"])
            page_extent = max((other["y"] + other["height"] for other in items), default=1)
            page_width = max((other["x"] + other["width"] for other in items), default=1)
            repeated_top_label = key in repeated_header_keys and item["y"] <= page_extent * 0.15
            pretitle_callout = bool(
                title_position
                and title_position[0] == page_number
                and item["y"] < title_position[1]["y"]
                and item["width"] <= page_width * 0.15
            )
            is_noise = _is_shape_page_label(item["text"]) or repeated_top_label or pretitle_callout
            blocks.append(context.block(
                "paragraph", item["text"],
                location=Location(page=page_number, bbox=[item["x"], item["y"], item["x"] + item["width"], item["y"] + item["height"]]),
                heading_path=[],
                content_origin="native_text",
                required=not is_noise,
                critical=False,
                status="skipped" if is_noise else "success",
                metadata={
                    "drawingml_anchor_index": item["anchor_index"],
                    **({"skip_reason": "repeated_header_page_label_or_pretitle_callout"} if is_noise else {}),
                },
            ))

    image_occurrences: list[dict[str, Any]] = []
    for page_number, items in enumerate(pages, start=1):
        for item in items:
            for occurrence in item.get("image_occurrences", []):
                image_occurrences.append({
                    **occurrence,
                    "page": page_number,
                    "element_index": item["anchor_index"],
                })
    image_blocks: list = []
    _append_media_occurrences(
        path,
        context,
        image_blocks,
        image_occurrences,
        heading_path=[],
        seen_payloads=set(),
        starting_index=0,
    )
    blocks.extend(image_blocks)

    if not any(block.type == "heading" and block.metadata.get("semantic_role") == "document_title" for block in blocks):
        blocks.append(context.block(
            "heading", title,
            heading_path=[title],
            content_origin="derived_normalization",
            required=True,
            critical=True,
            verbatim=False,
            metadata={
                "level": 1,
                "semantic_role": "document_title",
                "title_source": "filename",
                "title_from_filename": True,
            },
        ))
    blocks = _drawingml_sort_blocks(blocks)
    _annotate_reading_order(blocks, title)
    layout_risks = _docx_layout_risks(path)
    title_semantics_status = "reliable" if title_source != "filename" else "inferred"
    context.provenance.derivation_chain.append({
        "operation": "ooxml_drawingml_extraction",
        "source": str(path),
        "page_count": page_count,
    })
    return DocumentIR(
        source_id=context.provenance.source_id,
        title=title,
        provenance=context.provenance,
        blocks=blocks,
        metadata={
            "adapter": "docx_adapter",
            "fallback": "drawingml_shape_parser",
            "page_count": page_count,
            "drawingml_anchor_count": sum(len(page) for page in pages),
            "visual_occurrence_count": len(image_occurrences),
            "document_title": title,
            "title_source": title_source,
            "layout_semantics_status": "needs_review" if layout_risks else "reliable",
            "document_title_semantics_status": title_semantics_status,
        },
        warnings=[
            "docx_drawingml_rendering_caution",
            *layout_risks,
            *(["document_title_inferred_from_filename"] if title_semantics_status != "reliable" else []),
        ],
    )


def parse(path: Path, context: AdapterContext) -> DocumentIR:
    native_text_length = _native_text_length(path)
    shape_text_length = _shape_text_count(path)
    if native_text_length >= 100 or shape_text_length < 100:
        native_document = _native_parse(path, context)
        native_document.metadata.update({"native_text_length": native_text_length, "shape_text_length": shape_text_length})
        return native_document
    try:
        shape_document = _drawingml_parse(path, context)
        shape_document.metadata.update({"native_text_length": native_text_length, "shape_text_length": shape_text_length})
        return shape_document
    except Exception as shape_exc:
        try:
            converted = _convert_to_pdf(path, context.work_dir)
            context.provenance.derivation_chain.append({
                "operation": "libreoffice_docx_to_pdf",
                "source": str(path),
                "derived": str(converted),
            })
            derived = parse_pdf_document(
                converted,
                context,
                content_origin_override="derived_normalization",
                derived_from=str(path),
            )
            derived.metadata.update({
                "adapter": "docx_adapter",
                "fallback": "libreoffice_pdf_render",
                "native_text_length": native_text_length,
                "shape_text_length": shape_text_length,
                "drawingml_error": str(shape_exc),
            })
            return derived
        except Exception as exc:
            native_document = _native_parse(path, context)
            native_document.warnings.extend([f"drawingml_fallback_failed:{shape_exc}", f"rendered_pdf_fallback_failed:{exc}"])
            native_document.metadata.update({
                "native_text_length": native_text_length,
                "shape_text_length": shape_text_length,
                "fallback": "native_ooxml_only",
            })
            return native_document
